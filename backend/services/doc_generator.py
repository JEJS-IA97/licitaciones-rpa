import openpyxl
from docx import Document
from pathlib import Path
from config.settings import PLANTILLAS_DIR, LICITACIONES_DATA_DIR

def reemplazar_placeholders_word(doc: Document, reemplazos: dict):
    """Reemplaza variables tipo {{VARIABLE}} en párrafos y tablas de Word."""
    for p in doc.paragraphs:
        for key, val in reemplazos.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in reemplazos.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(val))

def generar_propuesta_word(id_licitacion: str, datos_empresa: dict, datos_ia: dict, nombre_salida: str):
    """Genera un archivo Word a partir de la plantilla base reemplazando datos."""
    plantilla_path = PLANTILLAS_DIR / "coimsa" / "plantilla_declaracion.docx"
    if not plantilla_path.exists():
        # Si no existe la plantilla específica, crea un documento base estándar
        doc = Document()
        doc.add_heading(f"DECLARACIÓN JURADA Y ACEPTACIÓN - {datos_empresa.get('razon_social')}", 0)
    else:
        doc = Document(str(plantilla_path))

    reemplazos = {
        "{{RAZON_SOCIAL}}": datos_empresa.get("razon_social", "COIMSA C.A."),
        "{{RUT_EMPRESA}}": datos_empresa.get("rut", "12.345.678-9"),
        "{{REPRESENTANTE}}": datos_empresa.get("representante", "Representante Legal"),
        "{{ID_LICITACION}}": id_licitacion,
        "{{CARGO_LICITACION}}": datos_ia.get("cargo_licitacion", "Servicio Generales")
    }

    reemplazar_placeholders_word(doc, reemplazos)

    dir_salida = LICITACIONES_DATA_DIR / id_licitacion / "archivos_finales"
    dir_salida.mkdir(parents=True, exist_ok=True)
    
    path_salida = dir_salida / f"{nombre_salida}.docx"
    doc.save(str(path_salida))
    return path_salida

def generar_excel_economico(id_licitacion: str, datos_ia: dict, nombre_salida: str):
    """Llena una planilla de Excel con la estructura presupuestaria solicitada."""
    path_salida = LICITACIONES_DATA_DIR / id_licitacion / "archivos_finales" / f"{nombre_salida}.xlsx"
    path_salida.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Propuesta Económica"

    ws.append(["Cargo / Insumo", "Cantidad", "Costo Unitario", "Subtotal"])

    total = 0
    for personal in datos_ia.get("personal_solicitado", []):
        cargo = personal.get("cargo", "Personal")
        cant = personal.get("cantidad", 1)
        sueldo = personal.get("remuneracion_minima", 0)
        subtotal = cant * sueldo
        total += subtotal
        ws.append([cargo, cant, sueldo, subtotal])

    ws.append([])
    ws.append(["TOTAL PRESUPUESTO", "", "", total])

    wb.save(str(path_salida))
    return path_salida